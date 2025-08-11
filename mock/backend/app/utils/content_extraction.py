"""Content extraction utilities for the AI Legal Assistant.

This module provides text extraction from various document formats including
PDF, DOCX, and plain text files with proper error handling and encoding detection.
"""

import logging
from pathlib import Path
from typing import Optional

import aiofiles
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class ContentExtractionError(Exception):
    """Base exception for content extraction errors."""
    pass


class UnsupportedFormatError(ContentExtractionError):
    """Exception raised when file format is not supported for extraction."""
    pass


class CorruptedFileError(ContentExtractionError):
    """Exception raised when file is corrupted or unreadable."""
    pass


async def extract_text_content(file_path: Path, content_type: str) -> str:
    """Extract text content from a file based on its content type.
    
    Args:
        file_path: Path to the file to extract content from
        content_type: MIME type of the file
        
    Returns:
        str: Extracted text content
        
    Raises:
        ContentExtractionError: If extraction fails
        UnsupportedFormatError: If file format is not supported
        CorruptedFileError: If file is corrupted or unreadable
    """
    try:
        content_type = content_type.lower().strip()
        
        logger.info(f"Extracting text from {file_path} (type: {content_type})")
        
        if content_type == "text/plain":
            return await _extract_text_from_txt(file_path)
        elif content_type == "application/pdf":
            return await _extract_text_from_pdf(file_path)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await _extract_text_from_docx(file_path)
        else:
            raise UnsupportedFormatError(f"Text extraction not supported for {content_type}")
            
    except (ContentExtractionError, UnsupportedFormatError, CorruptedFileError):
        raise
    except Exception as e:
        logger.error(f"Unexpected error extracting text from {file_path}: {e}")
        raise ContentExtractionError(f"Text extraction failed: {e}") from e


async def _extract_text_from_txt(file_path: Path) -> str:
    """Extract text from a plain text file with encoding detection.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        str: Text content
        
    Raises:
        CorruptedFileError: If file cannot be read or decoded
    """
    try:
        # Try common encodings in order of preference
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                    content = await f.read()
                    
                logger.info(f"Successfully read text file with {encoding} encoding")
                return content.strip()
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Failed to read with {encoding}: {e}")
                continue
        
        # If all encodings failed, try reading as binary and decode with error handling
        try:
            async with aiofiles.open(file_path, 'rb') as f:
                raw_data = await f.read()
            
            # Try to decode with error replacement
            content = raw_data.decode('utf-8', errors='replace')
            logger.warning(f"Text file decoded with error replacement")
            return content.strip()
            
        except Exception as e:
            raise CorruptedFileError(f"Cannot read text file: {e}") from e
            
    except CorruptedFileError:
        raise
    except Exception as e:
        logger.error(f"Error extracting text from TXT file: {e}")
        raise CorruptedFileError(f"Text file extraction failed: {e}") from e


async def _extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        str: Extracted text content
        
    Raises:
        CorruptedFileError: If PDF is corrupted or unreadable
    """
    try:
        # Read PDF file
        async with aiofiles.open(file_path, 'rb') as f:
            pdf_data = await f.read()
        
        # Create PDF reader from bytes
        try:
            from io import BytesIO
            pdf_reader = PdfReader(BytesIO(pdf_data))
        except Exception as e:
            raise CorruptedFileError(f"Cannot read PDF file: {e}") from e
        
        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            logger.warning(f"PDF file {file_path} is encrypted, attempting to decrypt")
            try:
                # Try to decrypt with empty password
                pdf_reader.decrypt("")
            except Exception as e:
                raise CorruptedFileError(f"Cannot decrypt PDF file: {e}") from e
        
        # Extract text from all pages
        text_content = []
        total_pages = len(pdf_reader.pages)
        
        if total_pages == 0:
            raise CorruptedFileError("PDF file contains no pages")
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                else:
                    logger.warning(f"No text found on page {page_num + 1} of PDF")
            except Exception as e:
                logger.warning(f"Error extracting text from PDF page {page_num + 1}: {e}")
                continue
        
        if not text_content:
            logger.warning(f"No text content extracted from PDF {file_path}")
            return ""
        
        # Join all page content
        full_text = "\n\n".join(text_content).strip()
        
        logger.info(f"Extracted text from PDF: {total_pages} pages, {len(full_text)} characters")
        return full_text
        
    except CorruptedFileError:
        raise
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise CorruptedFileError(f"PDF extraction failed: {e}") from e


async def _extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        str: Extracted text content
        
    Raises:
        CorruptedFileError: If DOCX is corrupted or unreadable
    """
    try:
        # Load DOCX document
        try:
            # python-docx doesn't support async directly, so we read the file first
            async with aiofiles.open(file_path, 'rb') as f:
                docx_data = await f.read()
            
            from io import BytesIO
            document = DocxDocument(BytesIO(docx_data))
            
        except Exception as e:
            raise CorruptedFileError(f"Cannot read DOCX file: {e}") from e
        
        # Extract text from paragraphs
        text_content = []
        
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        if not text_content:
            logger.warning(f"No text content found in DOCX {file_path}")
            return ""
        
        # Join all content
        full_text = "\n".join(text_content).strip()
        
        logger.info(f"Extracted text from DOCX: {len(full_text)} characters")
        return full_text
        
    except CorruptedFileError:
        raise
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise CorruptedFileError(f"DOCX extraction failed: {e}") from e


def extract_text_content_sync(file_path: Path, content_type: str) -> str:
    """Synchronous version of extract_text_content for non-async contexts.
    
    Args:
        file_path: Path to the file to extract content from
        content_type: MIME type of the file
        
    Returns:
        str: Extracted text content
        
    Raises:
        ContentExtractionError: If extraction fails
    """
    try:
        content_type = content_type.lower().strip()
        
        logger.info(f"Extracting text from {file_path} (type: {content_type}) - sync")
        
        if content_type == "text/plain":
            return _extract_text_from_txt_sync(file_path)
        elif content_type == "application/pdf":
            return _extract_text_from_pdf_sync(file_path)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return _extract_text_from_docx_sync(file_path)
        else:
            raise UnsupportedFormatError(f"Text extraction not supported for {content_type}")
            
    except (ContentExtractionError, UnsupportedFormatError, CorruptedFileError):
        raise
    except Exception as e:
        logger.error(f"Unexpected error extracting text from {file_path}: {e}")
        raise ContentExtractionError(f"Text extraction failed: {e}") from e


def _extract_text_from_txt_sync(file_path: Path) -> str:
    """Synchronous text extraction from TXT files."""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
        except Exception:
            continue
    
    # Fallback with error replacement
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        return raw_data.decode('utf-8', errors='replace').strip()
    except Exception as e:
        raise CorruptedFileError(f"Cannot read text file: {e}") from e


def _extract_text_from_pdf_sync(file_path: Path) -> str:
    """Synchronous text extraction from PDF files."""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PdfReader(f)
        
        if pdf_reader.is_encrypted:
            pdf_reader.decrypt("")
        
        text_content = []
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            except Exception as e:
                logger.warning(f"Error extracting text from PDF page: {e}")
                continue
        
        return "\n\n".join(text_content).strip()
        
    except Exception as e:
        raise CorruptedFileError(f"PDF extraction failed: {e}") from e


def _extract_text_from_docx_sync(file_path: Path) -> str:
    """Synchronous text extraction from DOCX files."""
    try:
        document = DocxDocument(file_path)
        
        text_content = []
        
        # Extract from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content).strip()
        
    except Exception as e:
        raise CorruptedFileError(f"DOCX extraction failed: {e}") from e


async def get_document_summary(file_path: Path, content_type: str) -> dict:
    """Get a summary of document content including metadata and preview.
    
    Args:
        file_path: Path to the document file
        content_type: MIME type of the file
        
    Returns:
        dict: Document summary with metadata and text preview
    """
    try:
        # Extract full text
        full_text = await extract_text_content(file_path, content_type)
        
        # Basic text analysis
        word_count = len(full_text.split()) if full_text else 0
        char_count = len(full_text)
        line_count = len(full_text.splitlines()) if full_text else 0
        
        # Create preview (first 500 characters)
        preview = full_text[:500] + "..." if len(full_text) > 500 else full_text
        
        return {
            "content_type": content_type,
            "file_size_bytes": file_path.stat().st_size,
            "extraction_successful": True,
            "text_length": char_count,
            "word_count": word_count,
            "line_count": line_count,
            "preview": preview,
            "has_content": bool(full_text.strip()),
            "error": None
        }
        
    except Exception as e:
        logger.error(f"Error creating document summary: {e}")
        return {
            "content_type": content_type,
            "file_size_bytes": file_path.stat().st_size if file_path.exists() else 0,
            "extraction_successful": False,
            "text_length": 0,
            "word_count": 0,
            "line_count": 0,
            "preview": "",
            "has_content": False,
            "error": str(e)
        }
